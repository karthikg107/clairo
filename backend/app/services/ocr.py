"""
OCR service — Google Cloud Vision primary, AWS Textract fallback.

SECURITY (non-negotiable):
- Document bytes purged from memory IMMEDIATELY after OCR completes
- NEVER log document content — only safe metadata (page count, confidence stats)
- All numbers flagged as always-verify regardless of confidence score
- No file ever written to disk during OCR

Architecture:
- Images (JPG, PNG, HEIC): Google Cloud Vision DOCUMENT_TEXT_DETECTION
- PDF/DOCX: Direct text extraction (confidence assumed 100% for digital text)
- On GCV failure: automatic retry with AWS Textract
"""
from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass, field
from enum import Enum

from app.core.logging import get_logger

logger = get_logger(__name__)


class OcrUnavailableError(Exception):
    """
    OCR could not read the file — e.g. an image OCR engine is unavailable,
    or the photo was unreadable. Surfaced to the user as a clear
    "couldn't read that image, try a PDF" message rather than a raw 500.
    """


# HEIC/HEIF (default iPhone photo format) needs an extra Pillow plugin.
# Registered once at import — best-effort so a missing plugin never breaks
# non-HEIC paths.
try:
    import pillow_heif  # type: ignore[import]

    pillow_heif.register_heif_opener()
except Exception:  # noqa: S110 — HEIC support is optional
    pass

# ── Types ─────────────────────────────────────────────────────────────────────

class ConfidenceLevel(str, Enum):
    HIGH   = "high"    # > 80%
    MEDIUM = "medium"  # 50–80%
    LOW    = "low"     # < 50%
    NUMBER = "number"  # Always verify — regardless of confidence


@dataclass
class OcrWord:
    text: str
    confidence: float          # 0.0–1.0 (1.0 for digital text extraction)
    confidence_level: ConfidenceLevel
    bounding_box: dict | None = None   # {x, y, w, h} in normalised coords


@dataclass
class OcrPage:
    page_number: int
    words: list[OcrWord] = field(default_factory=list)

    @property
    def text(self) -> str:
        return " ".join(w.text for w in self.words)

    @property
    def low_confidence_ratio(self) -> float:
        if not self.words:
            return 0.0
        low = sum(1 for w in self.words if w.confidence_level == ConfidenceLevel.LOW)
        return low / len(self.words)


@dataclass
class OcrResult:
    pages: list[OcrPage] = field(default_factory=list)
    source: str = "unknown"    # "gcv" | "textract" | "direct"
    total_pages: int = 0

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)


# ── Helpers ───────────────────────────────────────────────────────────────────

_NUMBER_PATTERN = re.compile(
    r"""
    (?:
        \d{1,3}(?:[.,]\d{3})*(?:[.,]\d+)?  # 1,234.56
        | \d+(?:[.,]\d+)?                   # 42 or 3.14
        | [€$£¥₹₩]\s*\d                    # currency prefixes
        | \d\s*[%°]                         # percentages / degrees
    )
    """,
    re.VERBOSE,
)


def _classify_word(text: str, confidence: float) -> ConfidenceLevel:
    """
    Classify a word by confidence level.
    ALL numbers are flagged as always-verify regardless of OCR confidence.
    """
    if _NUMBER_PATTERN.search(text):
        return ConfidenceLevel.NUMBER
    if confidence > 0.80:
        return ConfidenceLevel.HIGH
    if confidence >= 0.50:
        return ConfidenceLevel.MEDIUM
    return ConfidenceLevel.LOW


# ── Google Cloud Vision ───────────────────────────────────────────────────────

async def _ocr_with_gcv(image_bytes: bytes, mime_type: str) -> OcrResult:
    """
    OCR via Google Cloud Vision DOCUMENT_TEXT_DETECTION.
    Returns per-word confidence scores.

    SECURITY: image_bytes is consumed here and should be deleted by caller immediately.
    """
    from google.cloud import vision  # type: ignore[import]

    client = vision.ImageAnnotatorClient()
    image = vision.Image(content=image_bytes)
    features = [vision.Feature(type_=vision.Feature.Type.DOCUMENT_TEXT_DETECTION)]

    response = client.annotate_image({"image": image, "features": features})

    if response.error.message:
        raise RuntimeError(f"GCV error: {response.error.message}")

    pages: list[OcrPage] = []
    for page_idx, page in enumerate(response.full_text_annotation.pages):
        words: list[OcrWord] = []
        for block in page.blocks:
            for paragraph in block.paragraphs:
                for word in paragraph.words:
                    text = "".join(sym.text for sym in word.symbols)
                    conf = word.confidence if word.confidence else 0.5
                    # Bounding box (normalised)
                    verts = word.bounding_box.vertices
                    bb = {
                        "x": verts[0].x / page.width if page.width else 0,
                        "y": verts[0].y / page.height if page.height else 0,
                        "w": (verts[1].x - verts[0].x) / page.width if page.width else 0,
                        "h": (verts[2].y - verts[0].y) / page.height if page.height else 0,
                    }
                    words.append(OcrWord(
                        text=text,
                        confidence=conf,
                        confidence_level=_classify_word(text, conf),
                        bounding_box=bb,
                    ))
        pages.append(OcrPage(page_number=page_idx + 1, words=words))

    result = OcrResult(pages=pages, source="gcv", total_pages=len(pages))
    logger.info(
        "ocr.gcv_complete",
        page_count=result.total_pages,
        # NEVER log text content
    )
    return result


# ── AWS Textract (fallback) ───────────────────────────────────────────────────

async def _ocr_with_textract(image_bytes: bytes) -> OcrResult:
    """
    AWS Textract fallback. Returns words with confidence scores.

    SECURITY: image_bytes consumed here, deleted by caller.
    """
    import asyncio

    import boto3  # type: ignore[import]

    loop = asyncio.get_event_loop()
    client = boto3.client("textract", region_name="us-east-1")

    response = await loop.run_in_executor(
        None,
        lambda: client.detect_document_text(Document={"Bytes": image_bytes}),
    )

    words: list[OcrWord] = []
    for block in response.get("Blocks", []):
        if block["BlockType"] == "WORD":
            text = block.get("Text", "")
            conf = block.get("Confidence", 50.0) / 100.0  # Textract gives 0–100
            words.append(OcrWord(
                text=text,
                confidence=conf,
                confidence_level=_classify_word(text, conf),
            ))

    result = OcrResult(
        pages=[OcrPage(page_number=1, words=words)],
        source="textract",
        total_pages=1,
    )
    logger.info("ocr.textract_complete", word_count=len(words))
    return result


# ── Tesseract (free, in-container — default image engine) ────────────────────

async def _ocr_with_tesseract(image_bytes: bytes) -> OcrResult:
    """
    OCR via the bundled Tesseract engine (no external account / cost). Runs
    in a thread — Tesseract is CPU-bound and blocking.

    SECURITY: image_bytes consumed here, deleted by caller.
    """
    import asyncio

    import pytesseract  # type: ignore[import]
    from PIL import Image  # type: ignore[import]

    def _run() -> list[OcrWord]:
        image = Image.open(io.BytesIO(image_bytes))
        if image.mode not in ("L", "RGB"):
            image = image.convert("RGB")
        data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
        words: list[OcrWord] = []
        for text, raw_conf in zip(data["text"], data["conf"], strict=False):
            text = (text or "").strip()
            if not text:
                continue
            try:
                conf = float(raw_conf) / 100.0
            except (TypeError, ValueError):
                conf = 0.5
            if conf < 0:  # Tesseract uses -1 for non-text regions
                conf = 0.5
            words.append(
                OcrWord(text=text, confidence=conf, confidence_level=_classify_word(text, conf))
            )
        return words

    loop = asyncio.get_event_loop()
    words = await loop.run_in_executor(None, _run)

    result = OcrResult(
        pages=[OcrPage(page_number=1, words=words)],
        source="tesseract",
        total_pages=1,
    )
    logger.info("ocr.tesseract_complete", word_count=len(words))
    return result


# ── Direct text extraction (PDF / DOCX) ──────────────────────────────────────

def _extract_pdf_text(data: bytes) -> OcrResult:
    """Extract text directly from PDF. Confidence = 1.0 (digital text)."""
    import pypdf  # type: ignore[import]

    reader = pypdf.PdfReader(io.BytesIO(data))
    pages: list[OcrPage] = []

    for idx, page in enumerate(reader.pages):
        raw = page.extract_text() or ""
        words = [
            OcrWord(
                text=w,
                confidence=1.0,
                confidence_level=_classify_word(w, 1.0),
            )
            for w in raw.split()
            if w.strip()
        ]
        pages.append(OcrPage(page_number=idx + 1, words=words))

    result = OcrResult(pages=pages, source="direct", total_pages=len(pages))
    logger.info("ocr.pdf_extraction_complete", page_count=result.total_pages)
    return result


def _extract_docx_text(data: bytes) -> OcrResult:
    """Extract text directly from DOCX. Confidence = 1.0."""
    import docx  # type: ignore[import]

    doc = docx.Document(io.BytesIO(data))
    words: list[OcrWord] = []

    for para in doc.paragraphs:
        for w in para.text.split():
            if w.strip():
                words.append(OcrWord(
                    text=w,
                    confidence=1.0,
                    confidence_level=_classify_word(w, 1.0),
                ))

    result = OcrResult(
        pages=[OcrPage(page_number=1, words=words)],
        source="direct",
        total_pages=1,
    )
    logger.info("ocr.docx_extraction_complete", word_count=len(words))
    return result


# ── Public API ────────────────────────────────────────────────────────────────

async def run_ocr(
    data: bytes,
    mime_type: str,
    filename: str,
) -> OcrResult:
    """
    Run OCR on file bytes. Route:
    - PDF  → direct text extraction
    - DOCX → direct text extraction
    - Images → Google Cloud Vision → Textract fallback

    SECURITY: `data` must be deleted by the caller IMMEDIATELY after this returns.
    The caller is responsible for purging the bytes from memory.
    """
    from app.services.file_validation import AllowedMime

    logger.info("ocr.start", mime_type=mime_type)

    try:
        if mime_type == AllowedMime.PDF.value:
            return _extract_pdf_text(data)

        if mime_type == AllowedMime.DOCX.value:
            return _extract_docx_text(data)

        # ── Image path — pluggable engine (OCR_PROVIDER) ──────────────────
        # Default "tesseract": free, in-container, no external account.
        # Set OCR_PROVIDER=gcv (+ Google credentials) to switch to Google
        # Cloud Vision (higher accuracy) with a Textract fallback — no code
        # change needed, just config.
        provider = os.getenv("OCR_PROVIDER", "tesseract").strip().lower()
        try:
            if provider == "gcv":
                try:
                    return await _ocr_with_gcv(data, mime_type)
                except Exception as gcv_err:
                    logger.warning("ocr.gcv_failed_falling_back_to_textract", error=str(gcv_err))
                    return await _ocr_with_textract(data)
            return await _ocr_with_tesseract(data)
        except OcrUnavailableError:
            raise
        except Exception as exc:
            # Any image-OCR failure → a clear, user-facing error instead of a 500.
            logger.error("ocr.image_failed", provider=provider, error=str(exc))
            raise OcrUnavailableError(
                "Could not read text from that image. Try a clearer photo, "
                "or upload a PDF or Word document."
            ) from exc

    finally:
        # NOTE: caller must also `del data` after calling this function.
        # The `finally` block here logs completion but cannot del the caller's reference.
        logger.info("ocr.complete")
